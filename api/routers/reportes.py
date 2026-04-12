from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List
import io
import datetime

# PDF Libraries
from reportlab.lib.pagesizes import letter, landscape, portrait
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Excel Libraries
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill

# Word Libraries
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter(prefix="/reportes", tags=["reportes"])

class ReportData(BaseModel):
    title: str
    headers: List[str]
    rows: List[List[str]]
    filename: str = "reporte"

@router.post("/pdf")
async def generate_pdf(data: ReportData):
    buffer = io.BytesIO()
    
    # Orientation logic
    page_size = landscape(letter) if len(data.headers) > 5 else portrait(letter)
    doc = SimpleDocTemplate(buffer, pagesize=page_size)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'MainTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        textColor=colors.HexColor('#041329')
    )
    
    elements.append(Paragraph("MACUIN INDUSTRIAL", title_style))
    elements.append(Paragraph(data.title, styles['Heading2']))
    elements.append(Paragraph(f"Generado: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    elements.append(Paragraph(f"Total registros: {len(data.rows)}", styles['Normal']))
    elements.append(Spacer(1, 12))
    
    # Table logic
    table_data = [data.headers] + data.rows
    t = Table(table_data, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#041329')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#FFB3B1')),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
    ]))
    elements.append(t)
    
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={data.filename}.pdf"}
    )

@router.post("/xlsx")
async def generate_xlsx(data: ReportData):
    wb = Workbook()
    ws = wb.active
    ws.title = "Reporte"
    
    # Header styling
    header_font = Font(bold=True, color="FFB3B1")
    header_fill = PatternFill(start_color="041329", end_color="041329", fill_type="solid")
    
    ws.append(["MACUIN INDUSTRIAL"])
    ws.append([data.title])
    ws.append([f"Generado: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"])
    ws.append([]) # spacer
    
    ws.append(data.headers)
    header_row = 5
    for cell in ws[header_row]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    
    for row in data.rows:
        ws.append(row)
        
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={data.filename}.xlsx"}
    )

@router.post("/docx")
async def generate_docx(data: ReportData):
    doc = Document()
    
    # Title
    title = doc.add_heading('MACUIN INDUSTRIAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(data.title)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generado: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total registros: {len(data.rows)}")
    
    # Table
    table = doc.add_table(rows=1, cols=len(data.headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(data.headers):
        hdr_cells[i].text = header
        # Styling headers via XML but keeping it simple for now
        
    for row_data in data.rows:
        row_cells = table.add_row().cells
        for i, cell_val in enumerate(row_data):
            row_cells[i].text = str(cell_val)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={data.filename}.docx"}
    )
